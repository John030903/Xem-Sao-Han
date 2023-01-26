import streamlit as st
import datetime
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import time 
import json
import streamlit.components.v1 as components

# Tên các sao 
Nu = ("Thủy Diệu",
        "Kế Đô",
        "Vân Hớn",
        "Mộc Đức",
        "Thái Âm",
        "Thổ Tú",
        "La Hầu",
        "Thái Dương",
        "Thái Bạch"
    )
Nam = ("Mộc Đức",
        "La Hầu",
        "Thổ Tú",
        "Thủy Diệu",
        "Thái Bạch",
        "Thái Dương",
        "Vân Hớn",
        "Kế Đô",
        "Thái Âm"
    )
#Lấy năm
date = datetime.datetime.now()
namHienTai = date.year
# Tính ra sao hạn
def TinhSao(namSinh,gioiTinh):
    tuoi = namHienTai - namSinh
    sao = tuoi%9
    if "a" in gioiTinh:
        return str(Nam[sao])
    else:
        return str(Nu[sao])
st.set_page_config(
     page_title="Xem sao hạn năm " + str(namHienTai),
     page_icon="https://th.bing.com/th/id/R.aff394564778f0fd028ce35348a23384?rik=dI5iBsZkHdMrEA&riu=http%3a%2f%2fwww.downloadclipart.net%2flarge%2f3d-gold-star-png-pic.png&ehk=zF2KjUQ%2bGM6mjyUsyWh14LXUlQfBeOfJ01J2G5zkrJc%3d&risl=&pid=ImgRaw&r=0",
     layout="wide",
     initial_sidebar_state="expanded",
     menu_items={
         'Get Help': 'https://www.facebook.com/profile.php?id=100026754553136',
         'Report a bug': "https://forms.gle/hffqeLuazJWbnt5H7",
         'About': "# App dùng để tính sao và tạo ra một file word chứa danh sách người đã tính sao"
     }
 )
def local_css(file_name):
    with open(file_name) as f:
        st.markdown(f'<style>{f.read()}</style>', unsafe_allow_html=True)

def remote_css(url):
    st.markdown(f'<link href="{url}" rel="stylesheet">', unsafe_allow_html=True)  
     
local_css("style.css")
remote_css('https://fonts.googleapis.com/icon?family=Material+Icons')

st.title("XEM SAO HẠN NĂM " + str(namHienTai))
background = st.container()

tab1, tab2, tab3 = st.tabs(["Xem sao","Nhập danh sách", "File Excel.xlsx"])
with tab1:
    birthday, gender,emp = st.columns((2,2,3))
    with birthday:
        year = st.text_input('Năm sinh',placeholder='199x')
    with gender:
        gender = st.radio('Giới tính',['Nam','Nữ'])
    st.button("Xem sao hạn",key="Xem")
    if st.session_state.Xem:
        if year == "":
            st.error("Vui lòng nhập năm sinh")
        else:
            with open('YNghiaSao.json',encoding='utf-8') as f:
                Data = json.load(f)
                Sao = TinhSao(namSinh=int(year),gioiTinh=str(gender))
                txt = st.text_area("", "Sao " + Sao +"\n" +"\n"+ Data[Sao]
, height=300)
            components.html("""<a href="https://bienthuy.com/tam-linh/y-nghia-cac-sao-cung-sao-va-giai-han/">Nguồn: bienthuy.com</a>""",height=5800,)
with tab2:
    st.header("Nhập danh sách")
    form = st.form("my_form",clear_on_submit=True)
    with form.container():
        name, birthday, gender = st.columns((4,1,1))
        with name:
            name = st.text_input('Họ và tên',placeholder='Nguyễn Văn A')
        with birthday:
            year = st.text_input('Năm sinh',placeholder='199x')
        with gender:
            gender = st.radio('Giới tính',['Nam','Nữ'])

        if "count" not in st.session_state:
            st.session_state.count = 0
        add = form.form_submit_button("Thêm vào danh sách")
        if add:
            st.session_state.count += 1
            st.session_state['member' + str(st.session_state.count)] = [name,year,gender]
    members = []
    for i in range(1,st.session_state.count+1):
        members.append(st.session_state['member'+str(i)])
    df = pd.DataFrame(data=members,columns=['Họ và tên','Năm sinh','Giới tính'],index= (i for i in range(1,st.session_state.count+1)))
    if st.session_state.count != 0:
        st.subheader("Danh sách đã nhập")
        st.table(df)
    buttonDone, down, emp = st.columns((2,2,6))
    buttonDone.button("Hoàn thành",key="Done")
    if st.session_state.Done:
        df.to_excel("Data.xlsx",index=False)
        with open("Data.xlsx", "rb") as fileData:
            down.download_button("Tải file .xlsx dùng cho năm sau",data=fileData,file_name="Danh_Sach.xlsx", key="downData")

with open("Example.xlsx", "rb") as example:
  tab3.download_button("Tải file .xlsx mẫu",data=example,file_name="Danh_Sach.xlsx", key="downloaded")
if st.session_state.downloaded:
    tab3.success('✔️ Đã tải xuống')
tab3.header("Tải lên danh sách người cần tính sao (.xlsx)")
file = tab3.file_uploader(label="", type="xlsx",key = "uploaded")

if st.session_state.uploaded or st.session_state.Done:
    time.sleep(2)
    FONT_PARAGRAPH = 12
    

    # Tạo document
    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = 'Time News Roman'
    font.size = Pt(FONT_PARAGRAPH)
    
    # Thêm phần heading
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = p.add_run("SAO HẠN NĂM "+ str(namHienTai))
    p.font.size = Pt(18)
    p.font.bold = True
    # Tạo bảng
    table = doc.add_table(rows=1,cols=2)
    table.style = 'Table Grid'
    # Thêm 2 ô đầu tiên trong bảng
    row = table.rows[0].cells
    p = row[0].add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = p.add_run("Họ và Tên")
    p.font.size = Pt(FONT_PARAGRAPH)
    p.font.bold = True
    p = row[1].add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = p.add_run("Sao Hạn")
    p.font.size = Pt(FONT_PARAGRAPH)
    p.font.bold = True
    # Đọc file danh sách người cần tính sao 
    if st.session_state.uploaded:
        read_file = pd.read_excel (file, header=0)
    else:
        read_file = pd.read_excel ("Data.xlsx", header=0)
    for i in range(read_file.shape[0]):
        saoHan = TinhSao(read_file["Năm sinh"][i],read_file["Giới tính"][i])
        row = table.add_row().cells
        row[0].add_paragraph(read_file["Họ và tên"][i])
        p=row[1].add_paragraph(saoHan)
        p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    # Lưu file
    with st.spinner('Vui lòng chờ...'):
        time.sleep(1)
    doc.save("Sao Hạn Năm "+str(namHienTai)+".docx")

    with st.empty():
        st.success('✔️ Đã tính xong')
        time.sleep(0.5)
        with open("Sao Hạn Năm "+str(namHienTai)+".docx", "rb") as file:
            st.download_button(label="Tải file đã tính sao .docx", data=file, file_name="Sao Hạn Năm "+str(namHienTai)+".docx")
empty = background.empty()
i = 0
while True:
    if i == 3:
        i = 0
    empty.image("background"+str(i)+".jpg")
    i = i+1
    time.sleep(3)
